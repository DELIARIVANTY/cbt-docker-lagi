"""
Script untuk membuat dokumen Word lengkap: Normalisasi + Rancangan Program
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_complete_document():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('RANCANGAN BASIS DATA DAN PROGRAM', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_heading('SISTEM COMPUTER BASED TEST (CBT)', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ===== BAGIAN 4.2: NORMALISASI (simplified version for Word) =====
    doc.add_page_break()
    doc.add_heading('4.2. Rancangan Basis Data', 1)
    doc.add_heading('4.2.1. Normalisasi', 2)
    
    intro = doc.add_paragraph(
        'Proses normalisasi basis data Sistem CBT dimulai dari bentuk tidak normal (UNF) '
        'hingga Boyce-Codd Normal Form (BCNF). Berikut adalah tahapan normalisasi yang dilakukan:'
    )
    
    # UNF
    doc.add_heading('A. Unnormalized Form (UNF)', 3)
    doc.add_paragraph(
        'Pada tahap awal, seluruh data sistem CBT berada dalam satu tabel besar dengan banyak '
        'repeating groups (kelompok data berulang) yang menyebabkan redundansi tinggi.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Masalah UNF:').bold = True
    doc.add_paragraph('Terdapat repeating groups', style='List Bullet')
    doc.add_paragraph('Redundansi data tinggi', style='List Bullet')
    doc.add_paragraph('Anomali insert, update, delete', style='List Bullet')
    
    # 1NF
    doc.add_heading('B. First Normal Form (1NF)', 3)
    doc.add_paragraph(
        'Memecah repeating groups menjadi tabel-tabel terpisah. Setiap atribut bernilai atomik '
        'dan setiap tabel memiliki primary key.'
    )
    doc.add_paragraph(
        'Tabel yang dihasilkan: CustomUser, Jurusan, Kelas, MataPelajaran, BankSoal, ButirSoal, '
        'Ujian, SesiUjian, JawabanSiswa, JadwalPengawas.'
    )
    
    # 2NF
    doc.add_heading('C. Second Normal Form (2NF)', 3)
    doc.add_paragraph(
        'Menghilangkan partial dependency. Semua atribut non-key harus fully functionally dependent '
        'terhadap seluruh primary key. Semua tabel yang ada sudah memenuhi 2NF.'
    )
    
    # 3NF
    doc.add_heading('D. Third Normal Form (3NF)', 3)
    doc.add_paragraph(
        'Menghilangkan transitive dependency. Tidak ada atribut non-key yang bergantung pada '
        'atribut non-key lainnya. Semua tabel sudah memenuhi 3NF.'
    )
    
    # BCNF
    doc.add_heading('E. Boyce-Codd Normal Form (BCNF)', 3)
    doc.add_paragraph(
        'Untuk setiap functional dependency X → Y, X harus berupa superkey. '
        'Database CBT sudah mencapai BCNF, bentuk normal tertinggi yang praktis untuk OLTP systems.'
    )
    
    # Class Diagram Note
    doc.add_heading('F. Class Diagram', 3)
    doc.add_paragraph(
        'Class diagram menggambarkan 12 tabel utama dalam 3 domain:'
    )
    doc.add_paragraph('Academic: Jurusan, Kelas, MataPelajaran', style='List Bullet')
    doc.add_paragraph('Accounts: CustomUser (multi-role)', style='List Bullet')
    doc.add_paragraph('Exams: BankSoal, ButirSoal, Ujian, SesiUjian, JawabanSiswa, JadwalPengawas', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('CATATAN: ').bold = True
    p.add_run('Class diagram visual dapat dilihat dalam file markdown atau dibuat menggunakan '
              'tools diagram seperti draw.io atau dbdiagram.io.')
    
    # ===== BAGIAN 4.2.2: SPESIFIKASI BASIS DATA (10 tabel utama) =====
    doc.add_page_break()
    doc.add_heading('4.2.2. Spesifikasi Basis Data', 2)
    
    tables_spec = [
        {
            'no': 1,
            'name': 'CustomUser',
            'file': 'accounts_customuser',
            'akronim': 'USR',
            'fungsi': 'Menyimpan data pengguna sistem dengan berbagai role (Admin, Guru, Siswa, Proktor, Waka)',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '650 karakter',
            'kunci': 'id (PK), username (Unique), nisn (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'username', 'Varchar', '150', 'Username login, Unique, Not Null'),
                ('3', 'password', 'Varchar', '128', 'Password terenkripsi, Not Null'),
                ('4', 'nama', 'Varchar', '255', 'Nama lengkap, Not Null'),
                ('5', 'role', 'Varchar', '10', 'Role: admin/guru/siswa/proktor/waka'),
                ('6', 'nisn', 'Varchar', '20', 'NISN untuk siswa, Unique'),
                ('7', 'kelas_id', 'Integer', '11', 'FK ke Kelas'),
            ]
        },
        {
            'no': 2,
            'name': 'Jurusan',
            'file': 'academic_jurusan',
            'akronim': 'JRS',
            'fungsi': 'Menyimpan data jurusan yang tersedia di sekolah',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '75 karakter',
            'kunci': 'id (PK), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '50', 'Nama jurusan, Not Null'),
                ('3', 'kode', 'Varchar', '10', 'Kode jurusan, Unique, Not Null'),
            ]
        },
        {
            'no': 3,
            'name': 'Kelas',
            'file': 'academic_kelas',
            'akronim': 'KLS',
            'fungsi': 'Menyimpan data kelas dengan relasi ke jurusan',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '80 karakter',
            'kunci': 'id (PK)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '50', 'Nama kelas, Not Null'),
                ('3', 'tingkat', 'Integer', '11', 'Tingkat kelas (10/11/12)'),
                ('4', 'jurusan_id', 'Integer', '11', 'FK ke Jurusan, Not Null'),
            ]
        },
        {
            'no': 4,
            'name': 'MataPelajaran',
            'file': 'academic_matapelajaran',
            'akronim': 'MPL',
            'fungsi': 'Menyimpan data mata pelajaran yang diajarkan',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '135 karakter',
            'kunci': 'id (PK), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '100', 'Nama mata pelajaran, Not Null'),
                ('3', 'kode', 'Varchar', '20', 'Kode mapel, Unique, Not Null'),
            ]
        },
        {
            'no': 5,
            'name': 'BankSoal',
            'file': 'exams_banksoal',
            'akronim': 'BNK',
            'fungsi': 'Menyimpan bank soal yang dibuat guru per mata pelajaran',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '270 karakter',
            'kunci': 'id (PK), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'judul', 'Varchar', '200', 'Judul bank soal, Not Null'),
                ('3', 'kode', 'Varchar', '20', 'Kode unik (auto-generated), Unique'),
                ('4', 'mapel_id', 'Integer', '11', 'FK ke MataPelajaran, Not Null'),
                ('5', 'guru_id', 'Integer', '11', 'FK ke CustomUser (guru), Not Null'),
            ]
        },
        {
            'no': 6,
            'name': 'ButirSoal',
            'file': 'exams_butirsoal',
            'akronim': 'BTR',
            'fungsi': 'Menyimpan butir soal individual (PG atau Essay)',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '2100 karakter',
            'kunci': 'id (PK)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'bank_soal_id', 'Integer', '11', 'FK ke BankSoal, Not Null'),
                ('3', 'jenis_soal', 'Varchar', '10', 'Jenis: PG atau ESSAY'),
                ('4', 'pertanyaan', 'Text', '-', 'Teks pertanyaan, Not Null'),
                ('5', 'opsi_a hingga e', 'Text', '-', 'Opsi jawaban untuk PG'),
                ('6', 'kunci_jawaban', 'Varchar', '1', 'Kunci jawaban (A/B/C/D/E)'),
                ('7', 'bobot', 'Integer', '11', 'Bobot nilai soal, Default: 1'),
            ]
        },
        {
            'no': 7,
            'name': 'Ujian',
            'file': 'exams_ujian',
            'akronim': 'UJN',
            'fungsi': 'Menyimpan data ujian menggunakan bank soal tertentu',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '280 karakter',
            'kunci': 'id (PK)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'bank_soal_id', 'Integer', '11', 'FK ke BankSoal, Not Null'),
                ('3', 'nama_ujian', 'Varchar', '200', 'Nama ujian, Not Null'),
                ('4', 'waktu_mulai', 'DateTime', '-', 'Waktu mulai ujian'),
                ('5', 'durasi', 'Integer', '11', 'Durasi dalam menit'),
                ('6', 'token', 'Varchar', '6', 'Token akses (auto-generated)'),
                ('7', 'semester', 'Varchar', '10', 'Semester: Ganjil/Genap'),
            ]
        },
        {
            'no': 8,
            'name': 'SesiUjian',
            'file': 'exams_sesiujian',
            'akronim': 'SSI',
            'fungsi': 'Menyimpan sesi ujian individual per siswa',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '350 karakter',
            'kunci': 'id (PK)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'ujian_id', 'Integer', '11', 'FK ke Ujian, Not Null'),
                ('3', 'siswa_id', 'Integer', '11', 'FK ke CustomUser (siswa)'),
                ('4', 'waktu_mulai', 'DateTime', '-', 'Waktu mulai sesi, Auto'),
                ('5', 'is_finished', 'Boolean', '1', 'Status selesai, Default: False'),
                ('6', 'status', 'Varchar', '20', 'ONGOING/WAITING_GRADE/GRADED'),
                ('7', 'nilai', 'Float', '-', 'Nilai akhir (0-100)'),
            ]
        },
        {
            'no': 9,
            'name': 'JawabanSiswa',
            'file': 'exams_jawabansiswa',
            'akronim': 'JWB',
            'fungsi': 'Menyimpan jawaban siswa per soal dalam sesi ujian',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '1050 karakter',
            'kunci': '(sesi_id, soal_id) Composite Unique',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'sesi_id', 'Integer', '11', 'FK ke SesiUjian, Not Null'),
                ('3', 'soal_id', 'Integer', '11', 'FK ke ButirSoal, Not Null'),
                ('4', 'jawaban', 'Varchar', '1', 'Jawaban PG (A/B/C/D/E)'),
                ('5', 'jawaban_essay', 'Text', '-', 'Jawaban essay/uraian'),
                ('6', 'score', 'Float', '-', 'Score per jawaban'),
                ('7', 'ragu_ragu', 'Boolean', '1', 'Status ragu-ragu'),
            ]
        },
        {
            'no': 10,
            'name': 'JadwalPengawas',
            'file': 'exams_jadwalpengawas',
            'akronim': 'JPW',
            'fungsi': 'Menyimpan assignment pengawas per kelas dalam ujian',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '45 karakter',
            'kunci': '(ujian_id, kelas_id) Composite Unique',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'ujian_id', 'Integer', '11', 'FK ke Ujian, Not Null'),
                ('3', 'kelas_id', 'Integer', '11', 'FK ke Kelas, Not Null'),
                ('4', 'proktor_id', 'Integer', '11', 'FK ke CustomUser (guru)'),
            ]
        },
    ]
    
    # Add each table specification
    for spec in tables_spec:
        doc.add_heading(f"Tabel {spec['no']}: {spec['name']}", 3)
        
        # Specification table
        spec_table = doc.add_table(rows=8, cols=2)
        spec_table.style = 'Light List Accent 1'
        
        spec_table.rows[0].cells[0].text = 'Nama File'
        spec_table.rows[0].cells[1].text = spec['file']
        spec_table.rows[1].cells[0].text = 'Akronim'
        spec_table.rows[1].cells[1].text = spec['akronim']
        spec_table.rows[2].cells[0].text = 'Fungsi'
        spec_table.rows[2].cells[1].text = spec['fungsi']
        spec_table.rows[3].cells[0].text = 'Tipe'
        spec_table.rows[3].cells[1].text = spec['tipe']
        spec_table.rows[4].cells[0].text = 'Organisasi File'
        spec_table.rows[4].cells[1].text = spec['organisasi']
        spec_table.rows[5].cells[0].text = 'Media'
        spec_table.rows[5].cells[1].text = spec['media']
        spec_table.rows[6].cells[0].text = 'Panjang Record'
        spec_table.rows[6].cells[1].text = spec['panjang']
        spec_table.rows[7].cells[0].text = 'Field Kunci'
        spec_table.rows[7].cells[1].text = spec['kunci']
        
        # Field structure
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Struktur Field:').bold = True
        
        field_table = doc.add_table(rows=1, cols=5)
        field_table.style = 'Light Grid Accent 1'
        hdr = field_table.rows[0].cells
        hdr[0].text = 'No'
        hdr[1].text = 'Nama Field'
        hdr[2].text = 'Tipe Data'
        hdr[3].text = 'Lebar'
        hdr[4].text = 'Keterangan'
        
        for field in spec['fields']:
            row = field_table.add_row().cells
            row[0].text = field[0]
            row[1].text = field[1]
            row[2].text = field[2]
            row[3].text = field[3]
            row[4].text = field[4]
        
        doc.add_paragraph()
    
    # ===== BAGIAN 4.3: RANCANGAN PROGRAM =====
    doc.add_page_break()
    doc.add_heading('4.3. Rancangan Program', 1)
    
    intro_program = doc.add_paragraph(
        'Rancangan program menjelaskan struktur navigasi dan tampilan aplikasi Sistem CBT '
        'berdasarkan role pengguna. Sistem memiliki 5 role utama: Administrator, Guru, Siswa, '
        'Proktor/Panitia, dan Waka Kurikulum.'
    )
    
    # ===== 4.3.1 Struktur Navigasi =====
    doc.add_heading('4.3.1. Struktur Navigasi Sistem', 2)
    
    doc.add_paragraph(
        'Sistem CBT memiliki struktur navigasi berbasis role dengan hak akses yang berbeda '
        'untuk setiap pengguna. Berikut adalah menu dan fitur untuk masing-masing role:'
    )
    
    # A. Navigasi Administrator
    doc.add_heading('A. Navigasi Administrator', 3)
    
    p = doc.add_paragraph()
    p.add_run('Menu Administrator:').bold = True
    
    admin_menus = [
        ('Dashboard', 'Statistik sistem, quick actions, overview data'),
        ('Manajemen User', 'Tambah/Edit/Hapus user, Import Excel, Generate kartu peserta'),
        ('Manajemen Akademik', 'Kelola Jurusan, Kelas, Mata Pelajaran, Import siswa ke kelas'),
        ('Manajemen Ujian', 'Lihat semua bank soal, ujian, jadwal pengawas'),
        ('Monitoring & Laporan', 'Monitoring real-time, Laporan nilai, Analisis ujian'),
    ]
    
    table_admin = doc.add_table(rows=1, cols=2)
    table_admin.style = 'Light Grid Accent 1'
    hdr = table_admin.rows[0].cells
    hdr[0].text = 'Menu'
    hdr[1].text = 'Fitur'
    
    for menu, fitur in admin_menus:
        row = table_admin.add_row().cells
        row[0].text = menu
        row[1].text = fitur
    
    doc.add_paragraph()
    
    # B. Navigasi Guru
    doc.add_heading('B. Navigasi Guru', 3)
    
    p = doc.add_paragraph()
    p.add_run('Menu Guru:').bold = True
    
    guru_menus = [
        ('Dashboard', 'Bank soal, ujian yang dibuat, ujian aktif, perlu koreksi'),
        ('Bank Soal', 'Buat bank soal, Tambah/Edit soal (PG/Essay), Import soal Excel'),
        ('Ujian', 'Buat ujian (pilih bank, kelas, jadwal), Monitoring, Lihat hasil'),
        ('Koreksi', 'Koreksi essay, Beri nilai, Submit nilai akhir'),
    ]
    
    table_guru = doc.add_table(rows=1, cols=2)
    table_guru.style = 'Light Grid Accent 1'
    hdr = table_guru.rows[0].cells
    hdr[0].text = 'Menu'
    hdr[1].text = 'Fitur'
    
    for menu, fitur in guru_menus:
        row = table_guru.add_row().cells
        row[0].text = menu
        row[1].text = fitur
    
    doc.add_paragraph()
    
    # C. Navigasi Siswa
    doc.add_heading('C. Navigasi Siswa', 3)
    
    p = doc.add_paragraph()
    p.add_run('Menu Siswa:').bold = True
    
    siswa_menus = [
        ('Dashboard', 'Ujian tersedia, ujian berlangsung, riwayat ujian'),
        ('Ujian Tersedia', 'Lihat detail ujian, Konfirmasi ikut ujian, Input token'),
        ('Lembar Ujian', 'Timer countdown, Jawab soal, Navigasi nomor, Ragu-ragu, Submit'),
        ('Riwayat', 'Lihat hasil ujian, Detail jawaban, Nilai akhir'),
    ]
    
    table_siswa = doc.add_table(rows=1, cols=2)
    table_siswa.style = 'Light Grid Accent 1'
    hdr = table_siswa.rows[0].cells
    hdr[0].text = 'Menu'
    hdr[1].text = 'Fitur'
    
    for menu, fitur in siswa_menus:
        row = table_siswa.add_row().cells
        row[0].text = menu
        row[1].text = fitur
    
    doc.add_paragraph()
    
    # D. Navigasi Proktor
    doc.add_heading('D. Navigasi Proktor/Panitia', 3)
    
    p = doc.add_paragraph()
    p.add_run('Menu Proktor:').bold = True
    
    proktor_menus = [
        ('Dashboard', 'Ujian yang diawasi, statistik kehadiran'),
        ('Monitoring Real-time', 'Status siswa, Progress pengerjaan, Waktu tersisa, Alert'),
        ('Jadwal Pengawasan', 'Daftar ujian diawasi, Detail kelas, Waktu & tempat'),
        ('Laporan', 'Rekap kehadiran, Rekap nilai per kelas, Export laporan'),
    ]
    
    table_proktor = doc.add_table(rows=1, cols=2)
    table_proktor.style = 'Light Grid Accent 1'
    hdr = table_proktor.rows[0].cells
    hdr[0].text = 'Menu'
    hdr[1].text = 'Fitur'
    
    for menu, fitur in proktor_menus:
        row = table_proktor.add_row().cells
        row[0].text = menu
        row[1].text = fitur
    
    doc.add_paragraph()
    
    # E. Navigasi Waka
    doc.add_heading('E. Navigasi Waka Kurikulum', 3)
    
    p = doc.add_paragraph()
    p.add_run('Menu Waka Kurikulum:').bold = True
    
    waka_menus = [
        ('Dashboard', 'Overview sistem, Statistik ujian & nilai, Trend'),
        ('Laporan Akademik', 'Per mata pelajaran, Per kelas, Per guru, Filter semester'),
        ('Analisis & Statistik', 'Analisis nilai, Perbandingan kelas, Grafik, Export Excel/PDF'),
    ]
    
    table_waka = doc.add_table(rows=1, cols=2)
    table_waka.style = 'Light Grid Accent 1'
    hdr = table_waka.rows[0].cells
    hdr[0].text = 'Menu'
    hdr[1].text = 'Fitur'
    
    for menu, fitur in waka_menus:
        row = table_waka.add_row().cells
        row[0].text = menu
        row[1].text = fitur
    
    # ===== 4.3.2 Diagram HIPO =====
    doc.add_page_break()
    doc.add_heading('4.3.2. Diagram HIPO (Hierarchy plus Input-Process-Output)', 2)
    
    doc.add_paragraph(
        'Diagram HIPO menggambarkan hierarki modul sistem dan alur Input-Process-Output '
        'untuk setiap fungsi utama dalam Sistem CBT.'
    )
    
    # HIPO Level 0
    doc.add_heading('A. HIPO Level 0 - Sistem CBT Keseluruhan', 3)
    
    hipo0_table = doc.add_table(rows=4, cols=1)
    hipo0_table.style = 'Medium Shading 1 Accent 1'
    
    hipo0_table.rows[0].cells[0].text = 'SISTEM CBT (Computer Based Test)'
    
    p = hipo0_table.rows[1].cells[0].paragraphs[0]
    p.clear()
    run = p.add_run('INPUT:')
    run.bold = True
    p.add_run('\n• Credentials User (username, password)')
    p.add_run('\n• Data Master (user, kelas, jurusan, mata pelajaran)')
    p.add_run('\n• Data Soal (bank soal, butir soal)')
    p.add_run('\n• Jawaban Siswa')
    
    p = hipo0_table.rows[2].cells[0].paragraphs[0]
    p.clear()
    run = p.add_run('PROCESS:')
    run.bold = True
    p.add_run('\n1. Autentikasi & Otorisasi')
    p.add_run('\n2. Manajemen Data Master')
    p.add_run('\n3. Manajemen Bank Soal & Ujian')
    p.add_run('\n4. Pelaksanaan Ujian')
    p.add_run('\n5. Penilaian Otomatis & Manual')
    p.add_run('\n6. Monitoring & Pelaporan')
    
    p = hipo0_table.rows[3].cells[0].paragraphs[0]
    p.clear()
    run = p.add_run('OUTPUT:')
    run.bold = True
    p.add_run('\n• Dashboard per Role')
    p.add_run('\n• Hasil Ujian & Nilai')
    p.add_run('\n• Laporan & Analisis')
    p.add_run('\n• Monitoring Real-time')
    
    doc.add_paragraph()
    
    # HIPO Level 1 - Modul Pelaksanaan Ujian (Detailed)
    doc.add_heading('B. HIPO Level 1 - Modul Pelaksanaan Ujian (Modul Kritis)', 3)
    
    hipo1_data = [
        {
            'modul': '5.1 Validasi & Mulai Ujian',
            'input': '• ID Siswa\n• Token Ujian\n• Device ID',
            'process': '1. Validasi token\n2. Cek jadwal aktif\n3. Cek duplikasi sesi\n4. Buat sesi baru\n5. Load soal',
            'output': '• Sesi Ujian dibuat\n• Lembar Soal\n• Timer countdown start'
        },
        {
            'modul': '5.2 Menjawab Soal',
            'input': '• ID Sesi\n• ID Soal\n• Jawaban (PG/Essay)\n• Status Ragu-ragu',
            'process': '1. Validasi sesi aktif\n2. Simpan jawaban\n3. Update timestamp\n4. Sync database\n5. Update UI',
            'output': '• Jawaban tersimpan\n• Status update\n• Progress terupdate'
        },
        {
            'modul': '5.3 Navigasi Soal',
            'input': '• Nomor Soal\n• Action (Next/Prev)',
            'process': '1. Load soal\n2. Load jawaban\n3. Update UI\n4. Track progress',
            'output': '• Soal ditampilkan\n• Jawaban loaded\n• Progress indicator'
        },
        {
            'modul': '5.4 Timer & Auto Submit',
            'input': '• Waktu Tersisa\n• Status Ujian',
            'process': '1. Countdown timer\n2. Cek waktu habis\n3. Auto submit\n4. Hitung nilai PG',
            'output': '• Timer display\n• Ujian selesai\n• Nilai sementara'
        },
        {
            'modul': '5.5 Submit Ujian',
            'input': '• ID Sesi\n• Trigger Submit\n• Konfirmasi',
            'process': '1. Konfirmasi submit\n2. Finalisasi sesi\n3. Hitung nilai PG\n4. Update status',
            'output': '• Status: FINISHED\n• Nilai tersimpan\n• Redirect hasil'
        },
    ]
    
    for hipo in hipo1_data:
        doc.add_heading(hipo['modul'], 4)
        
        hipo_table = doc.add_table(rows=3, cols=1)
        hipo_table.style = 'Light Grid Accent 1'
        
        p = hipo_table.rows[0].cells[0].paragraphs[0]
        p.clear()
        run = p.add_run('INPUT:')
        run.bold = True
        p.add_run('\n' + hipo['input'])
        
        p = hipo_table.rows[1].cells[0].paragraphs[0]
        p.clear()
        run = p.add_run('PROCESS:')
        run.bold = True
        p.add_run('\n' + hipo['process'])
        
        p = hipo_table.rows[2].cells[0].paragraphs[0]
        p.clear()
        run = p.add_run('OUTPUT:')
        run.bold = True
        p.add_run('\n' + hipo['output'])
        
        doc.add_paragraph()
    
    # ===== 4.3.3 Wireframe =====
    doc.add_page_break()
    doc.add_heading('4.3.3. Wireframe Tampilan Utama', 2)
    
    doc.add_paragraph(
        'Berikut adalah representasi wireframe untuk tampilan-tampilan utama sistem:'
    )
    
    # Wireframe descriptions
    wireframes = [
        ('Halaman Login', 'Form login dengan username dan password, logo sekolah, tombol login'),
        ('Dashboard Siswa', 'Statistik ujian (tersedia, berlangsung, riwayat), daftar ujian tersedia dengan detail'),
        ('Konfirmasi Ujian', 'Detail ujian, form input token akses 6 digit, tombol mulai ujian'),
        ('Lembar Ujian', 'Timer countdown, navigasi nomor soal, pertanyaan, opsi jawaban, indikator status, progress bar'),
        ('Dashboard Guru', 'Statistik (bank soal, ujian, perlu koreksi), menu cepat, ujian aktif, bank soal terbaru'),
        ('Monitoring Proktor', 'Daftar siswa, status pengerjaan, progress bar per siswa, waktu tersisa, statistik kehadiran'),
    ]
    
    doc.add_heading('Komponen Wireframe:', 3)
    
    for title, desc in wireframes:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)
    
    p = doc.add_paragraph()
    p.add_run('CATATAN: ').bold = True
    p.add_run('Wireframe detail dalam bentuk visual dapat dilihat dalam file markdown atau '
              'dibuat menggunakan tools seperti Figma, Balsamiq, atau MockFlow.')
    
    # ===== 4.3.4 Teknologi =====
    doc.add_heading('4.3.4. Teknologi yang Digunakan', 2)
    
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    hdr = tech_table.rows[0].cells
    hdr[0].text = 'Komponen'
    hdr[1].text = 'Teknologi'
    hdr[2].text = 'Keterangan'
    
    technologies = [
        ('Backend Framework', 'Django 5.x', 'Python web framework'),
        ('Frontend', 'HTML5, CSS3, JavaScript', 'Antarmuka pengguna'),
        ('Database', 'SQLite/PostgreSQL/MySQL', 'Database relasional'),
        ('CSS Framework', 'Bootstrap 5', 'Responsive design'),
        ('Icons', 'Font Awesome', 'Icon library'),
        ('Charts', 'Chart.js', 'Visualisasi data'),
        ('Authentication', 'Django Auth System', 'Autentikasi built-in'),
        ('Export', 'ReportLab/openpyxl', 'Generate PDF & Excel'),
    ]
    
    for comp, tech, ket in technologies:
        row = tech_table.add_row().cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = ket
    
    # ===== Kesimpulan =====
    doc.add_page_break()
    doc.add_heading('4.3.5. Kesimpulan Rancangan Program', 2)
    
    doc.add_paragraph(
        'Rancangan program Sistem CBT ini telah dirancang dengan mempertimbangkan:'
    )
    
    doc.add_paragraph('User-Centered Design: Navigasi sesuai role dan hak akses', style='List Bullet')
    doc.add_paragraph('Modular Structure: Modul independen dan mudah dimaintain', style='List Bullet')
    doc.add_paragraph('Clear Workflow: Alur kerja jelas dari persiapan hingga pelaporan', style='List Bullet')
    doc.add_paragraph('Real-time Monitoring: Fitur monitoring untuk pengawasan ujian', style='List Bullet')
    doc.add_paragraph('Responsive Interface: Desain responsif dan user-friendly', style='List Bullet')
    doc.add_paragraph('Scalable Architecture: Dapat dikembangkan sesuai kebutuhan', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Struktur navigasi dan diagram HIPO memberikan gambaran komprehensif tentang '
        'interaksi pengguna dengan sistem pada setiap tahap proses ujian berbasis komputer.'
    )

    # ===== BAGIAN 4.5.3: HAK AKSES =====
    doc.add_page_break()
    doc.add_heading('4.5.3. Hak Akses', 2)

    doc.add_paragraph(
        'Implementasi hak akses dalam Sistem CBT ini menerapkan prinsip Least Privilege, '
        'di mana setiap pengguna hanya diberikan akses sesuai dengan kebutuhan tugas dan fungsinya. '
        'Hak akses dikelompokkan menjadi dua kategori utama:'
    )

    # a. Petugas yang Berwenang
    doc.add_heading('a. Petugas yang Berwenang', 3)
    doc.add_paragraph(
        'Petugas yang berwenang adalah pengguna yang terlibat langsung dalam operasional teknis, '
        'manajemen data, dan pelaksanaan kegiatan evaluasi sehari-hari.'
    )

    petugas_roles = [
        ('Administrator (Superuser)', 
         'Memiliki akses penuh (root access) terhadap seluruh fitur dan data. Bertanggung jawab atas manajemen master data, maintenance, dan konfigurasi sistem.'),
        ('Guru Mata Pelajaran', 
         'Akses terbatas pada konten akademik. Bertanggung jawab atas Bank Soal, jadwal ujian kelas yang diampu, dan koreksi nilai.'),
        ('Proktor / Panitia Ujian', 
         'Fokus pada pengawasan teknis. Bertanggung jawab monitoring real-time, reset login peserta, dan memastikan integritas ujian.'),
    ]

    for role_name, desc in petugas_roles:
        p = doc.add_paragraph()
        run = p.add_run(f'• {role_name}: ')
        run.bold = True
        p.add_run(desc)

    # b. Pimpinan
    doc.add_heading('b. Pimpinan', 3)
    doc.add_paragraph(
        'Pimpinan adalah pengguna yang memiliki hak akses untuk memantau kinerja sistem dan '
        'hasil evaluasi secara manajerial tanpa terlibat dalam operasional teknis harian.'
    )

    pimpinan_roles = [
        ('Waka Kurikulum / Kepala Sekolah',
         'Akses Read-Only terhadap laporan dan statistik. Bertanggung jawab monitoring kinerja global, evaluasi hasil ujian, dan pengambilan keputusan berbasis data.'),
    ]

    for role_name, desc in pimpinan_roles:
        p = doc.add_paragraph()
        run = p.add_run(f'• {role_name}: ')
        run.bold = True
        p.add_run(desc)

    # c. Siswa
    doc.add_heading('c. Siswa / Peserta Ujian', 3)
    doc.add_paragraph(
        'Siswa merupakan pengguna akhir yang menjadi objek evaluasi dalam sistem CBT. '
        'Hak akses siswa terbatas pada:'
    )

    siswa_roles = [
        ('Siswa',
         'Wewenang akses pelaksanaan ujian sesuai jadwal dan kelas. Bertanggung jawab mengerjakan ujian dengan jujur dan memastikan jawaban tersimpan.'),
    ]

    for role_name, desc in siswa_roles:
        p = doc.add_paragraph()
        run = p.add_run(f'• {role_name}: ')
        run.bold = True
        p.add_run(desc)
    
    # Save document
    output_path = r'C:\Users\USER\.gemini\antigravity\brain\b945bedc-f0f7-454f-9fcd-259d9f7e84fb\Dokumentasi_Lengkap_CBT.docx'
    doc.save(output_path)
    print(f"✓ Dokumen Word lengkap berhasil dibuat: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_complete_document()
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
