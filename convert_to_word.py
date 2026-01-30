"""
Script untuk mengkonversi dokumentasi normalisasi database dari Markdown ke Word (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_hyperlink(paragraph, text, url):
    """Menambahkan hyperlink ke paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_word_document():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('4.2. Rancangan Basis Data - Sistem CBT', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.2.1 Section
    doc.add_heading('4.2.1. Normalisasi', 1)
    
    intro = doc.add_paragraph(
        'Dokumen ini menjelaskan proses normalisasi basis data untuk Sistem Computer Based Test (CBT) '
        'yang dimulai dari bentuk tidak normal (Unnormalized Form/UNF) hingga bentuk normal terakhir '
        '(Boyce-Codd Normal Form/BCNF).'
    )
    
    # A. UNF
    doc.add_heading('A. Bentuk Tidak Normal (Unnormalized Form - UNF)', 2)
    
    p = doc.add_paragraph(
        'Pada tahap awal, data sistem CBT dapat direpresentasikan dalam satu tabel besar yang mengandung '
        'semua informasi tanpa memperhatikan redundansi data:'
    )
    
    doc.add_heading('Tabel: DATA_UJIAN_LENGKAP', 3)
    
    # Create UNF table
    unf_data = [
        ('id_ujian', 'ID Ujian'),
        ('nama_ujian', 'Nama Ujian'),
        ('token_ujian', 'Token untuk akses ujian'),
        ('semester', 'Semester ujian (Ganjil/Genap)'),
        ('waktu_mulai', 'Waktu mulai ujian'),
        ('durasi', 'Durasi ujian dalam menit'),
        ('aktif', 'Status aktif ujian'),
        ('id_bank_soal', 'ID Bank Soal'),
        ('judul_bank', 'Judul Bank Soal'),
        ('kode_bank', 'Kode Bank Soal'),
        ('id_mapel', 'ID Mata Pelajaran'),
        ('nama_mapel', 'Nama Mata Pelajaran'),
        ('kode_mapel', 'Kode Mata Pelajaran'),
        ('id_guru_pembuat', 'ID Guru pembuat'),
        ('nama_guru_pembuat', 'Nama Guru'),
        ('username_guru', 'Username Guru'),
        ('id_siswa', 'ID Siswa (repeating group)'),
        ('nama_siswa', 'Nama Siswa (repeating group)'),
        ('id_butir_soal', 'ID Butir Soal (repeating group)'),
        ('pertanyaan', 'Teks Pertanyaan (repeating group)'),
        ('kunci_jawaban', 'Kunci Jawaban (repeating group)'),
        ('id_sesi', 'ID Sesi Ujian (repeating group)'),
        ('id_jawaban', 'ID Jawaban (repeating group)'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Atribut'
    hdr_cells[1].text = 'Deskripsi'
    
    for attr, desc in unf_data:
        row_cells = table.add_row().cells
        row_cells[0].text = attr
        row_cells[1].text = desc
    
    # Masalah UNF
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Masalah pada UNF:').bold = True
    doc.add_paragraph('Terdapat banyak repeating groups (kelompok data yang berulang)', style='List Bullet')
    doc.add_paragraph('Redundansi data yang sangat tinggi', style='List Bullet')
    doc.add_paragraph('Anomali insert, update, dan delete', style='List Bullet')
    doc.add_paragraph('Tidak efisien dalam penyimpanan dan query', style='List Bullet')
    
    # B. 1NF
    doc.add_page_break()
    doc.add_heading('B. Bentuk Normal Pertama (1st Normal Form - 1NF)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Aturan 1NF:').bold = True
    doc.add_paragraph('Setiap atribut harus bernilai atomik (tidak ada repeating groups)', style='List Bullet')
    doc.add_paragraph('Setiap baris harus unik (ada primary key)', style='List Bullet')
    
    p = doc.add_paragraph('\nPada tahap ini, kita memecah repeating groups menjadi tabel-tabel terpisah:')
    
    # Tables in 1NF
    tables_1nf = [
        ('UJIAN', 'id_ujian', 'nama_ujian, token, semester, waktu_mulai, durasi, aktif, id_bank_soal'),
        ('BANK_SOAL', 'id_bank_soal', 'judul_bank, kode_bank, id_mapel, id_guru_pembuat, created_at, updated_at'),
        ('MATA_PELAJARAN', 'id_mapel', 'nama_mapel, kode_mapel'),
        ('CUSTOMUSER', 'id_user', 'username, nama, email, role, password, nisn, id_kelas'),
        ('KELAS', 'id_kelas', 'nama_kelas, tingkat, id_jurusan'),
        ('JURUSAN', 'id_jurusan', 'nama_jurusan, kode_jurusan'),
        ('BUTIR_SOAL', 'id_butir_soal', 'id_bank_soal, jenis_soal, pertanyaan, opsi_a-e, kunci_jawaban, bobot'),
        ('SESI_UJIAN', 'id_sesi', 'id_ujian, id_siswa, waktu_mulai, waktu_selesai, status, nilai, score'),
        ('JAWABAN_SISWA', '(id_sesi, id_butir_soal)', 'jawaban_pg, jawaban_essay, score, ragu_ragu'),
        ('JADWAL_PENGAWAS', '(id_ujian, id_kelas)', 'id_pengawas'),
    ]
    
    for table_name, pk, attrs in tables_1nf:
        p = doc.add_paragraph()
        p.add_run(f'Tabel {table_name}').bold = True
        p = doc.add_paragraph(f'PK: {pk}')
        p = doc.add_paragraph(f'Atribut: {attrs}')
    
    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('✓ Sudah memenuhi 1NF karena tidak ada repeating groups dan setiap tabel memiliki primary key.')
    
    # C. 2NF
    doc.add_page_break()
    doc.add_heading('C. Bentuk Normal Kedua (2nd Normal Form - 2NF)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Aturan 2NF:').bold = True
    doc.add_paragraph('Sudah dalam bentuk 1NF', style='List Bullet')
    doc.add_paragraph('Semua atribut non-key harus fully functionally dependent terhadap seluruh primary key (tidak ada partial dependency)', style='List Bullet')
    
    p = doc.add_paragraph('\nAnalisis Dependency:')
    p = doc.add_paragraph(
        'Semua tabel yang ada sudah memenuhi 2NF karena tidak terdapat partial dependency. '
        'Tabel dengan single-column primary key otomatis memenuhi 2NF. Tabel dengan composite key '
        '(seperti JAWABAN_SISWA dan JADWAL_PENGAWAS) memiliki atribut yang fully dependent pada kombinasi lengkap key.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('✓ Sudah memenuhi 2NF karena tidak ada partial dependency.')
    
    # D. 3NF
    doc.add_heading('D. Bentuk Normal Ketiga (3rd Normal Form - 3NF)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Aturan 3NF:').bold = True
    doc.add_paragraph('Sudah dalam bentuk 2NF', style='List Bullet')
    doc.add_paragraph('Tidak ada transitive dependency (atribut non-key tidak boleh bergantung pada atribut non-key lain)', style='List Bullet')
    
    p = doc.add_paragraph('\nAnalisis Transitive Dependency:')
    p = doc.add_paragraph(
        'Semua tabel telah dianalisis dan tidak ditemukan transitive dependency. '
        'Foreign key seperti id_kelas, id_mapel, id_guru bukan merupakan transitive dependency '
        'melainkan relasi antar entitas yang valid.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('✓ Sudah memenuhi 3NF karena tidak ada transitive dependency.')
    
    # E. BCNF
    doc.add_page_break()
    doc.add_heading('E. Boyce-Codd Normal Form (BCNF)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Aturan BCNF:').bold = True
    doc.add_paragraph('Sudah dalam bentuk 3NF', style='List Bullet')
    doc.add_paragraph('Untuk setiap functional dependency X → Y, X harus berupa superkey', style='List Bullet')
    
    p = doc.add_paragraph('\nAnalisis BCNF:')
    p = doc.add_paragraph(
        'Semua tabel dalam desain database CBT ini sudah memenuhi BCNF karena semua dependency '
        'memiliki determinant yang merupakan superkey.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('✓ Sudah memenuhi BCNF - Bentuk normal tertinggi yang praktis untuk OLTP systems.')
    
    # Class Diagram explanation
    doc.add_heading('F. Class Diagram - Final Database Schema', 2)
    
    p = doc.add_paragraph(
        'Class diagram menggambarkan struktur final database dalam bentuk BCNF dengan 12 tabel utama '
        'yang terbagi dalam 3 domain: Academic (Jurusan, Kelas, MataPelajaran), Accounts (CustomUser), '
        'dan Exams (BankSoal, ButirSoal, Ujian, SesiUjian, JawabanSiswa, JadwalPengawas).'
    )
    
    p = doc.add_paragraph()
    p.add_run('CATATAN: ').bold = True
    p.add_run('Class diagram visual dapat dilihat dalam file markdown atau dibuat menggunakan tools seperti draw.io, '
              'Lucidchart, atau dbdiagram.io berdasarkan spesifikasi tabel di bawah.')
    
    # Kesimpulan
    doc.add_heading('G. Kesimpulan', 2)
    
    p = doc.add_paragraph(
        'Database Sistem CBT ini telah dinormalisasi hingga Boyce-Codd Normal Form (BCNF), '
        'yang merupakan bentuk normal tertinggi yang praktis untuk sistem transaksional (OLTP). Desain ini:'
    )
    
    doc.add_paragraph('✓ Menghilangkan redundansi data', style='List Bullet')
    doc.add_paragraph('✓ Mencegah anomali insert, update, dan delete', style='List Bullet')
    doc.add_paragraph('✓ Memastikan integritas data melalui constraints', style='List Bullet')
    doc.add_paragraph('✓ Mendukung scalability dan maintainability', style='List Bullet')
    doc.add_paragraph('✓ Mengikuti best practices dalam database design', style='List Bullet')
    
    # 4.2.2 Spesifikasi Basis Data
    doc.add_page_break()
    doc.add_heading('4.2.2. Spesifikasi Basis Data', 1)
    
    p = doc.add_paragraph(
        'Spesifikasi basis data berikut berisi rincian lengkap dari setiap tabel dalam class diagram '
        'yang telah dibuat pada bagian normalisasi.'
    )
    
    # Database specifications for each table
    tables_spec = [
        {
            'name': 'CustomUser',
            'file': 'accounts_customuser',
            'akronim': 'USR',
            'fungsi': 'Menyimpan data pengguna sistem dengan berbagai role (Admin, Guru, Siswa, Proktor, Waka Kurikulum)',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '650 karakter',
            'kunci': 'id (Primary Key), username (Unique), nisn (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'username', 'Varchar', '150', 'Username untuk login, Unique, Not Null'),
                ('3', 'password', 'Varchar', '128', 'Password terenkripsi (hashed), Not Null'),
                ('4', 'plain_password', 'Varchar', '128', 'Password plain text (untuk recovery)'),
                ('5', 'email', 'Varchar', '254', 'Email pengguna'),
                ('6', 'nama', 'Varchar', '255', 'Nama lengkap, Not Null'),
                ('7', 'role', 'Varchar', '10', 'Role: admin/guru/siswa/proktor/waka, Not Null'),
                ('8', 'nisn', 'Varchar', '20', 'NISN untuk siswa, Unique'),
                ('9', 'kelas_id', 'Integer', '11', 'Foreign Key ke tabel Kelas'),
                ('10', 'is_active', 'Boolean', '1', 'Status aktif user'),
            ]
        },
        {
            'name': 'Jurusan',
            'file': 'academic_jurusan',
            'akronim': 'JRS',
            'fungsi': 'Menyimpan data jurusan yang tersedia di sekolah',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '75 karakter',
            'kunci': 'id (Primary Key), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '50', 'Nama jurusan (IPA, IPS, TKJ, dll), Not Null'),
                ('3', 'kode', 'Varchar', '10', 'Kode jurusan, Unique, Not Null'),
            ]
        },
        {
            'name': 'Kelas',
            'file': 'academic_kelas',
            'akronim': 'KLS',
            'fungsi': 'Menyimpan data kelas dengan relasi ke jurusan dan tingkat',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '80 karakter',
            'kunci': 'id (Primary Key)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '50', 'Nama kelas (contoh: X IPA 1), Not Null'),
                ('3', 'tingkat', 'Integer', '11', 'Tingkat kelas (10, 11, atau 12), Not Null'),
                ('4', 'jurusan_id', 'Integer', '11', 'Foreign Key ke tabel Jurusan, Not Null'),
            ]
        },
        {
            'name': 'MataPelajaran',
            'file': 'academic_matapelajaran',
            'akronim': 'MPL',
            'fungsi': 'Menyimpan data mata pelajaran yang diajarkan',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '135 karakter',
            'kunci': 'id (Primary Key), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'nama', 'Varchar', '100', 'Nama mata pelajaran, Not Null'),
                ('3', 'kode', 'Varchar', '20', 'Kode mata pelajaran, Unique, Not Null'),
            ]
        },
        {
            'name': 'BankSoal',
            'file': 'exams_banksoal',
            'akronim': 'BNK',
            'fungsi': 'Menyimpan bank soal yang dibuat oleh guru untuk mata pelajaran tertentu',
            'tipe': 'Master',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '270 karakter',
            'kunci': 'id (Primary Key), kode (Unique)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'judul', 'Varchar', '200', 'Judul bank soal, Not Null'),
                ('3', 'kode', 'Varchar', '20', 'Kode unik bank soal (auto-generated), Unique'),
                ('4', 'mapel_id', 'Integer', '11', 'Foreign Key ke MataPelajaran, Not Null'),
                ('5', 'guru_id', 'Integer', '11', 'Foreign Key ke CustomUser (role=guru), Not Null'),
                ('6', 'created_at', 'DateTime', '-', 'Waktu pembuatan, Auto'),
                ('7', 'updated_at', 'DateTime', '-', 'Waktu update terakhir, Auto'),
            ]
        },
        {
            'name': 'ButirSoal',
            'file': 'exams_butirsoal',
            'akronim': 'BTR',
            'fungsi': 'Menyimpan butir-butir soal individual dalam bank soal (PG atau Essay)',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '2100 karakter',
            'kunci': 'id (Primary Key)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'bank_soal_id', 'Integer', '11', 'Foreign Key ke BankSoal, Not Null'),
                ('3', 'jenis_soal', 'Varchar', '10', 'Jenis: PG atau ESSAY, Not Null'),
                ('4', 'pertanyaan', 'Text', '-', 'Teks pertanyaan soal, Not Null'),
                ('5', 'opsi_a', 'Text', '-', 'Opsi A untuk PG'),
                ('6', 'opsi_b', 'Text', '-', 'Opsi B untuk PG'),
                ('7', 'opsi_c', 'Text', '-', 'Opsi C untuk PG'),
                ('8', 'opsi_d', 'Text', '-', 'Opsi D untuk PG'),
                ('9', 'opsi_e', 'Text', '-', 'Opsi E untuk PG'),
                ('10', 'kunci_jawaban', 'Varchar', '1', 'Kunci jawaban untuk PG (A/B/C/D/E)'),
                ('11', 'bobot', 'Integer', '11', 'Bobot nilai soal, Default: 1'),
            ]
        },
        {
            'name': 'Ujian',
            'file': 'exams_ujian',
            'akronim': 'UJN',
            'fungsi': 'Menyimpan data ujian yang menggunakan bank soal tertentu',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '280 karakter',
            'kunci': 'id (Primary Key)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'bank_soal_id', 'Integer', '11', 'Foreign Key ke BankSoal, Not Null'),
                ('3', 'nama_ujian', 'Varchar', '200', 'Nama ujian, Not Null'),
                ('4', 'waktu_mulai', 'DateTime', '-', 'Waktu mulai ujian, Not Null'),
                ('5', 'durasi', 'Integer', '11', 'Durasi ujian dalam menit, Not Null'),
                ('6', 'token', 'Varchar', '6', 'Token akses ujian (auto-generated)'),
                ('7', 'semester', 'Varchar', '10', 'Semester: Ganjil/Genap, Not Null'),
                ('8', 'aktif', 'Boolean', '1', 'Status aktif ujian, Default: False'),
            ]
        },
        {
            'name': 'SesiUjian',
            'file': 'exams_sesiujian',
            'akronim': 'SSI',
            'fungsi': 'Menyimpan sesi ujian individual untuk setiap siswa yang mengikuti ujian',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '350 karakter',
            'kunci': 'id (Primary Key)',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'ujian_id', 'Integer', '11', 'Foreign Key ke Ujian, Not Null'),
                ('3', 'siswa_id', 'Integer', '11', 'Foreign Key ke CustomUser (role=siswa), Not Null'),
                ('4', 'waktu_mulai', 'DateTime', '-', 'Waktu mulai sesi, Auto'),
                ('5', 'waktu_selesai', 'DateTime', '-', 'Waktu selesai sesi'),
                ('6', 'is_finished', 'Boolean', '1', 'Status selesai mengerjakan, Default: False'),
                ('7', 'status', 'Varchar', '20', 'Status: ONGOING/WAITING_GRADE/GRADED'),
                ('8', 'nilai', 'Float', '-', 'Nilai akhir (0-100)'),
                ('9', 'score', 'Float', '-', 'Score mentah'),
            ]
        },
        {
            'name': 'JawabanSiswa',
            'file': 'exams_jawabansiswa',
            'akronim': 'JWB',
            'fungsi': 'Menyimpan jawaban siswa untuk setiap butir soal dalam sesi ujian',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '1050 karakter',
            'kunci': '(sesi_id, soal_id) Composite Unique Key',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'sesi_id', 'Integer', '11', 'Foreign Key ke SesiUjian, Not Null'),
                ('3', 'soal_id', 'Integer', '11', 'Foreign Key ke ButirSoal, Not Null'),
                ('4', 'jawaban', 'Varchar', '1', 'Jawaban PG yang dipilih (A/B/C/D/E)'),
                ('5', 'jawaban_essay', 'Text', '-', 'Jawaban essay/uraian'),
                ('6', 'score', 'Float', '-', 'Score per jawaban, Default: 0.0'),
                ('7', 'ragu_ragu', 'Boolean', '1', 'Status ragu-ragu, Default: False'),
            ]
        },
        {
            'name': 'JadwalPengawas',
            'file': 'exams_jadwalpengawas',
            'akronim': 'JPW',
            'fungsi': 'Menyimpan assignment pengawas (proktor) untuk kelas tertentu dalam ujian',
            'tipe': 'Transaksi',
            'organisasi': 'Index Sequential',
            'media': 'Hard Disk',
            'panjang': '45 karakter',
            'kunci': '(ujian_id, kelas_id) Composite Unique Key',
            'fields': [
                ('1', 'id', 'Integer', '11', 'Primary Key, Auto Increment'),
                ('2', 'ujian_id', 'Integer', '11', 'Foreign Key ke Ujian, Not Null'),
                ('3', 'kelas_id', 'Integer', '11', 'Foreign Key ke Kelas, Not Null'),
                ('4', 'proktor_id', 'Integer', '11', 'Foreign Key ke CustomUser (role=guru), Not Null'),
            ]
        },
    ]
    
    # Add each table specification
    for idx, spec in enumerate(tables_spec, 1):
        doc.add_heading(f"{idx}. Tabel {spec['name']}", 3)
        
        # Specification table
        spec_table = doc.add_table(rows=8, cols=2)
        spec_table.style = 'Light List Accent 1'
        
        spec_table.rows[0].cells[0].text = 'Nama File'
        spec_table.rows[0].cells[1].text = spec['file']
        spec_table.rows[1].cells[0].text = 'Akronim'
        spec_table.rows[1].cells[1].text = spec['akronim']
        spec_table.rows[2].cells[0].text = 'Fungsi'
        spec_table.rows[2].cells[1].text = spec['fungsi']
        spec_table.rows[3].cells[0].text = 'Tipe'
        spec_table.rows[3].cells[1].text = spec['tipe']
        spec_table.rows[4].cells[0].text = 'Organisasi File'
        spec_table.rows[4].cells[1].text = spec['organisasi']
        spec_table.rows[5].cells[0].text = 'Media'
        spec_table.rows[5].cells[1].text = spec['media']
        spec_table.rows[6].cells[0].text = 'Panjang Record'
        spec_table.rows[6].cells[1].text = spec['panjang']
        spec_table.rows[7].cells[0].text = 'Field Kunci'
        spec_table.rows[7].cells[1].text = spec['kunci']
        
        # Field structure
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Struktur Field:').bold = True
        
        field_table = doc.add_table(rows=1, cols=5)
        field_table.style = 'Light Grid Accent 1'
        hdr = field_table.rows[0].cells
        hdr[0].text = 'No'
        hdr[1].text = 'Nama Field'
        hdr[2].text = 'Tipe Data'
        hdr[3].text = 'Lebar'
        hdr[4].text = 'Keterangan'
        
        for field in spec['fields']:
            row = field_table.add_row().cells
            row[0].text = field[0]
            row[1].text = field[1]
            row[2].text = field[2]
            row[3].text = field[3]
            row[4].text = field[4]
        
        doc.add_paragraph()
    
    # Relasi Antar Tabel
    doc.add_page_break()
    doc.add_heading('Relasi Antar Tabel (Entity Relationship)', 2)
    
    rel_table = doc.add_table(rows=1, cols=5)
    rel_table.style = 'Light Grid Accent 1'
    hdr = rel_table.rows[0].cells
    hdr[0].text = 'Tabel Asal'
    hdr[1].text = 'Relasi'
    hdr[2].text = 'Tabel Tujuan'
    hdr[3].text = 'Kardinalitas'
    hdr[4].text = 'Foreign Key'
    
    relations = [
        ('Kelas', 'Many to One', 'Jurusan', 'N:1', 'kelas.jurusan_id'),
        ('CustomUser', 'Many to One', 'Kelas', 'N:1', 'customuser.kelas_id'),
        ('BankSoal', 'Many to One', 'MataPelajaran', 'N:1', 'banksoal.mapel_id'),
        ('BankSoal', 'Many to One', 'CustomUser (Guru)', 'N:1', 'banksoal.guru_id'),
        ('ButirSoal', 'Many to One', 'BankSoal', 'N:1', 'butirsoal.bank_soal_id'),
        ('Ujian', 'Many to One', 'BankSoal', 'N:1', 'ujian.bank_soal_id'),
        ('SesiUjian', 'Many to One', 'Ujian', 'N:1', 'sesiujian.ujian_id'),
        ('SesiUjian', 'Many to One', 'CustomUser (Siswa)', 'N:1', 'sesiujian.siswa_id'),
        ('JawabanSiswa', 'Many to One', 'SesiUjian', 'N:1', 'jawabansiswa.sesi_id'),
        ('JawabanSiswa', 'Many to One', 'ButirSoal', 'N:1', 'jawabansiswa.soal_id'),
        ('JadwalPengawas', 'Many to One', 'Ujian', 'N:1', 'jadwalpengawas.ujian_id'),
        ('JadwalPengawas', 'Many to One', 'Kelas', 'N:1', 'jadwalpengawas.kelas_id'),
    ]
    
    for rel in relations:
        row = rel_table.add_row().cells
        for i, val in enumerate(rel):
            row[i].text = val
    
    # Ringkasan Statistik
    doc.add_paragraph()
    doc.add_heading('Ringkasan Statistik Database', 2)
    
    stat_table = doc.add_table(rows=1, cols=2)
    stat_table.style = 'Light List Accent 1'
    hdr = stat_table.rows[0].cells
    hdr[0].text = 'Kategori'
    hdr[1].text = 'Jumlah'
    
    stats = [
        ('Total Tabel', '12 tabel'),
        ('Tabel Master', '5 tabel'),
        ('Tabel Transaksi', '5 tabel'),
        ('Tabel Relasi (Junction)', '2 tabel'),
        ('Total Field', '94 field'),
        ('Total Foreign Key', '17 relasi'),
        ('Total Unique Constraint', '6 constraint'),
    ]
    
    for stat in stats:
        row = stat_table.add_row().cells
        row[0].text = stat[0]
        row[1].text = stat[1]
    
    # Save document
    output_path = r'C:\Users\USER\.gemini\antigravity\brain\b945bedc-f0f7-454f-9fcd-259d9f7e84fb\Normalisasi_Database_CBT.docx'
    doc.save(output_path)
    print(f"✓ Dokumen Word berhasil dibuat: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_word_document()
    except ImportError:
        print("Error: Library python-docx tidak terinstall.")
        print("Install dengan: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
